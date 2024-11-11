import glob
from string import punctuation, digits
import os
import re
import docx
from docx.shared import Inches
from nltk.corpus import stopwords
import nltk.data
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import nltk

# Ensure NLTK data is downloaded
nltk.download('punkt_tab')
nltk.download('stopwords')

class TextProcessor:
    def __init__(self):
        # stopwords and tokenizer loaded here
        self.s = set(stopwords.words('english'))
        self.tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')
        # empty dictionary to be used to store values read from file
        self.dictionary = {}

    def read_file(self, filename):
        """reads a file and puts them into sentences"""
        with open(f'{filename}', encoding="utf8") as f:
            # reading file as a whole
            document = f.read()
            # splitting into sentences
            sentences = self.tokenizer.tokenize(document)
            # goes through each sentence one by one
            for text_inputs in sentences:
                self.append_words(text_inputs.lstrip(punctuation), filename)

    def append_words(self, text_input, filename):
        """takes sentence and filenames and appends them to dictionary"""
        original_text_input = text_input
        # digits and punctuation are removed, input is split
        text_input = text_input.translate(str.maketrans('', '', digits))
        text_input = text_input.translate(str.maketrans('', '', punctuation))
        text_input = text_input.split()
        # base name contains the filename without file extension
        base_name = os.path.splitext(os.path.basename(filename))[0]
        # looking through each word in text
        for word in text_input:
            original_word = word
            # case lowered
            word = word.lower()

            # regex find all words in sentence
            outputs = (re.findall(fr'([^.]*{word}[^.]*)', original_text_input, re.IGNORECASE))
            # stripping out bits
            outputs = [x.strip('\n').strip("") for x in outputs]
            # joining together
            outputs = '\n'.join([str(item) for item in outputs])
            if word in self.dictionary:
                # adds to the word count for the word
                self.dictionary[word]["Word(Total Occurrences)"] += 1
                # check if document number has been added yet
                if base_name not in self.dictionary[word]["Documents"]:
                    self.dictionary[word]["Documents"].append(base_name)
                # check if the sentence is in the dictionary yet
                if outputs not in self.dictionary[word]["Sentences"]:
                    # if-else statements check for capitalization
                    # regex looks for word makes sure word in word is skipped
                    if original_word.istitle():
                        self.dictionary[word]["Sentences"].append(
                            str(re.sub(fr'\b{word}\b', "*" + word.capitalize() + '*', outputs, flags=re.I)))
                    else:
                        self.dictionary[word]["Sentences"].append(
                            str(re.sub(fr'\b{word}\b', "*" + word + '*', outputs, flags=re.I)))
            else:
                # create new entry if none found
                # makes sure to add the first count by default
                # highlights found word, makes sure to not find words within word
                # if-else statements check for capitalization
                if original_word.istitle():
                    self.dictionary[word] = {"Word(Total Occurrences)": 1, "Documents": [base_name],
                                             "Sentences": [
                                                 re.sub(fr'\b{word}\b', "*" + word.capitalize() + '*', outputs,
                                                        flags=re.I) + '  ']}
                else:
                    self.dictionary[word] = {"Word(Total Occurrences)": 1, "Documents": [base_name],
                                             "Sentences": [
                                                 re.sub(fr'\b{word}\b', "*" + word + '*', outputs, flags=re.I) + '   ']} 

    def counted_values(self):
        """Removing words that have a smaller count"""
        new_dict = {}
        for key in self.dictionary:
            # looking only for words with a count over 40
            if self.dictionary[key]["Word(Total Occurrences)"] > 40:
                if key not in self.s:
                    # these words are being removed as they are common english language words
                    new_dict[key] = self.dictionary[key]
        # return new dictionary with smaller count
        return new_dict


import matplotlib.pyplot as plt

class DocumentCreator:
    @staticmethod
    def create_document(new_dict):
        """Creates a document in word containing a table"""
        # creates an instance of document
        doc = docx.Document()

        # add a large heading
        doc.add_heading('Word Frequency Counter', 0)

        # creating a table object
        table = doc.add_table(rows=1, cols=2)
        # sets the style as a grid
        table.style = 'Table Grid'
        # allows adjustment of table width
        table.allow_autofit = False
        table.autofit = False

        # set the table headers
        row = table.rows[0].cells
        row[0].text = 'Word count and documents'
        row[1].text = 'Sentences'

        # adding rows for each word contained in new_dict dictionary
        for key, item in new_dict.items():
            row = table.add_row().cells
            # capitalize each Word and add number of occurrences and add document names
            row[0].text = f'{key.capitalize()}, Count: {new_dict[key]["Word(Total Occurrences)"]}\n\nText Sources:\n{", ".join(new_dict[key]["Documents"])} '
            # remove duplicates using list to dict to list
            new_sentences = list(dict.fromkeys(new_dict[key]["Sentences"]))
            # make sure full stops are in
            sentence = ".\n\n".join([item.strip(" ") for item in new_sentences]) + "."
            # dont append full stops when question mark "?" symbol
            result = re.sub(r'[?][.]', '?', sentence)
            row[1].text = result

        # adjust the widths to suitable size
        widths = (Inches(2), Inches(4))
        for row in table.rows:
            # loops through indexes using count
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # save the document
        doc.save('frequency.docx')

    @staticmethod
    def generate_word_cloud(new_dict):
        """Generates a word cloud image based on word frequency in new_dict"""
        # Extract word frequencies for word cloud
        word_frequencies = {key: item["Word(Total Occurrences)"] for key, item in new_dict.items()}

        # Generate word cloud
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate_from_frequencies(word_frequencies)

        # Save the word cloud image
        wordcloud.to_file("wordcloud.png")

        # Optionally display the word cloud
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis('off')
        plt.show()

    @staticmethod
    def generate_frequency_graph(new_dict, top_n=10):
        """Generates a bar graph for the most frequent words"""
        # Extract word frequencies for the graph
        word_frequencies = {key: item["Word(Total Occurrences)"] for key, item in new_dict.items()}
        
        # Sort the word frequencies in descending order and take the top N
        sorted_words = sorted(word_frequencies.items(), key=lambda x: x[1], reverse=True)[:top_n]
        
        # Prepare data for plotting
        words, counts = zip(*sorted_words)
        
        # Create the bar chart
        plt.figure(figsize=(10, 5))
        plt.bar(words, counts, color='skyblue')
        plt.xlabel('Words')
        plt.ylabel('Frequency')
        plt.title(f'Top {top_n} Most Frequent Words')
        plt.xticks(rotation=45, ha='right')
        
        # Save the graph
        plt.tight_layout()
        plt.savefig('frequency_graph.png')
        
        # Optionally display the graph
        plt.show()


def main():
    processor = TextProcessor()
    
    # opens each file
    for filename in glob.iglob('test_docs/*.txt'):
        processor.read_file(filename)

    new_dict = processor.counted_values()
    # passes dictionary to be used to create document
    DocumentCreator.create_document(new_dict)
    # Generate a word cloud based on word frequency
    DocumentCreator.generate_word_cloud(new_dict)
    # Generate a frequency graph for top 10 most frequent words
    DocumentCreator.generate_frequency_graph(new_dict)


# runs the script
if __name__ == '__main__':
    main()
